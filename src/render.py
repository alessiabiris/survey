#functions for displaying and analysing the survey output

from __future__ import annotations
from typing import Dict, Any, List
import pandas as pd

#loops through all sections, questions and flattens everything into a table
def extract_codebook(survey: Dict[str, Any]) -> pd.DataFrame:
    rows: List[dict] = []
    for sec in survey.get("sections", []):
        for q in sec.get("questions", []):
            rows.append({
                "question_id": q.get("id"),
                "section": sec.get("title"),
                "text": q.get("text"),
                "type": q.get("type"),
                "options": " | ".join(q.get("options") or []),
                "topic": q.get("topic"),
                "analysis_tag": q.get("analysis_tag"),
                "required": q.get("required", True),
            })
    return pd.DataFrame(rows)

#count questions across sections
def count_questions(survey: Dict[str, Any]) -> int:
    n = 0
    for sec in survey.get("sections", []):
        n += len(sec.get("questions", []))
    return n


from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK
from io import BytesIO

def generate_survey_docx(survey: dict) -> bytes:
    doc = Document()
    
    # Title
    title = doc.add_heading("Survey", level=0)
    
    for sec in survey.get("sections", []):
        # Section title
        doc.add_heading(sec.get("title", ""), level=1)
        
        for q in sec.get("questions", []):
            # Question text
            q_para = doc.add_paragraph()
            q_para.add_run(f"{q.get('id')} — ").bold = True
            q_para.add_run(q.get("text", ""))
            
            qtype = q.get("type")
            opts = q.get("options") or []
            
            if qtype in ("single_choice", "likert_5", "likert_7"):
                for opt in opts:
                    doc.add_paragraph(f"○  {opt}", style="List Bullet")
            
            elif qtype == "multi_choice":
                for opt in opts:
                    doc.add_paragraph(f"☐  {opt}", style="List Bullet")
            
            elif qtype == "free_text":
                doc.add_paragraph("_" * 50)
            
            elif qtype == "numeric":
                doc.add_paragraph("[Enter number: ______ ]")
            
            # Add spacing
            doc.add_paragraph()
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
